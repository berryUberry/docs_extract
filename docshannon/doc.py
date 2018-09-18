# encoding: utf-8 
""" 
@author: wangkai 
@contact: berryberryry@gmail.com
@version: 1.0 
@license: Apache Licence 
@file: doc.py
@time: 18-9-18 下午2:23 

这一行开始写关于本文件的说明与解释 
"""

import zipfile
import os

from docx import Document
from docshannon.table.table import Table, Cell


class DOCDocument(object):

    def __init__(self, doc_path, img_path=None):
        doc = Document(doc_path)
        self.paragraphs = doc.paragraphs
        self.tables = doc.tables
        self.doc_path = doc_path
        self.img_path = img_path

    def get_doc_text(self):
        for paragraph in self.paragraphs:
            print(paragraph.text)

    def get_doc_paragraphs(self):
        return self.paragraphs

    def get_doc_tables(self):
        for table in self.tables:
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(row_count):
                print("*"*40)
                for j in range(col_count):
                    # print(table.cell(i,j)._element._grid_col)
                    cell = table.cell(i, j)
                    print(cell.text, ',', cell._element.left, ',', cell._element.right, ',', cell._element.top, ',', cell._element.bottom)
                    if table.cell(i, j)._element._grid_col == j:
                        print(table.cell(i, j).text)

        return self.tables

    def get_doc_images(self):
        # read file
        self.zipf = zipfile.ZipFile(self.doc_path)
        self.filelist = self.zipf.namelist()

        if self.img_path:
            # extract images
            for fname in self.filelist:
                _, extension = os.path.splitext(fname)
                if extension in [".jpg", ".jpeg", ".png", ".bmp"]:
                    dst_fname = os.path.join(self.img_path, os.path.basename(fname))
                    with open(dst_fname, "wb") as dst_f:
                        dst_f.write(self.zipf.read(fname))

    def get_doc_tables_json(self):
        id = 0
        for table in self.tables:
            shannon_table = Table(id, len(table.rows), len(table.columns))
            id += 1

